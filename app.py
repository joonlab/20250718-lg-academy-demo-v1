import streamlit as st
import pandas as pd
import docx
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from io import BytesIO

# --- 페이지 설정 ---
st.set_page_config(
    page_title="성과관리 운영 현황 대시보드",
    layout="wide"
)

# --- 데이터 정의 (이미지 내용) ---
# 이 부분의 데이터를 수정하여 대시보드 내용을 변경할 수 있습니다.
title = "[작성요청] 성과관리 운영 현황 (예시적 / 각 사 상황에 맞게 기재)"

# 등급 배분
dist_method = "절대평가"
process_flow = "팀장 점수 평가/등급 제안 → 등급 심의 위원회 실시(위원장: 담당) → 개인별 등급 피드백(확정) → 이의 제기 → 최종 확정"

# 등급별 분포 현황 표 데이터
table_data = {
    '평가': ['S', 'A', 'B', 'C', 'D'],
    '책임(인원비중: 50%)': ['XX%', 'XX%', 'XX%', 'XX%', 'XX%'],
    '선임(인원비중: 30%)': ['XX%', 'XX%', 'XX%', 'XX%', 'XX%'],
    '사원(인원비중: 20%)': ['XX%', 'XX%', 'XX%', 'XX%', 'XX%'],
    'Total': ['XX%', 'XX%', 'XX%', 'XX%', 'XX%']
}
df = pd.DataFrame(table_data).set_index('평가').T

# 구성원 VOE
voe_list = [
    "🟢 절대평가 전환 이후 진급이나 연차에 따른 평가 왜곡은 많이 개선된 것으로 느껴짐",
    "🟢 절대평가 전환 이후 조직 내 경쟁이 줄어들어 동료 의식이 강화되고 팀워크에 도움이 되는 것 같음",
    "🟢 상대평가는 연초에 어떤 업무를 부여받는 지에 따라 평가가 결정되는 경우가 다수인데, 절대평가를 통해 이러한 부분이 많이 개선되었음",
    "🔴 절대평가 전환 이후 관대화가 많이 이뤄져 적당히 해도 B를 받는다고 느끼거나 B 평가를 수용하지 못함",
    "🔴 절대평가 전환 이후 재원은 한정된 상황에서 평가가 관대화되다 보니 전반적인 임금경쟁력이 낮아지는 경향이 있고, 고성과자에 대한 동기부여도 되지 않음. 차라리 상대평가가 나을 것 같음",
    "🔴 평가 기준이 조직별로 달라 타 팀의 A와 나 팀의 A가 같지 않은 경향이 있음"
]

# 평가 운영상의 Issue
issue_list = [
    "- Pay Band를 기준으로 평가/보상을 분리하지 않아 저 직급 인원의 저평가 현상이 나타남.",
    "- 평가 공정성 제고를 위해 수시 성과관리를 강화하려고 하나, 조직 책임자들의 업무 Load 증가로 인한 불만 증가",
    "- 성과 평가 보완을 위해 동료 평가를 도입하려고 하나, 구성원 수용성 부족",
    "- 역량과 성과를 보상에 연계하려고 하나, 역량 평가에 대한 공신력 부족",
    "- 이의제기 절차에 대한 평가 변경 / 유지에 대한 모호성",
    "- 보상과 평가 분리 요구",
    "(기타 각 사에서 성과관리 강화를 위해 개선이 필요한 사항들 / 타사 의견을 들어보고 싶은 사례들에 대해 기재)"
]


# --- Word 문서 생성 함수 (2단 레이아웃 적용) ---
def create_word_document():
    doc = docx.Document()
    
    # 1. 문서 제목
    doc.add_heading(title, level=1)
    doc.add_paragraph()

    # 2. 등급 배분 섹션 (이 섹션은 전체 너비 사용)
    container_box = doc.add_table(rows=1, cols=1).cell(0,0)
    container_box.text = '' # 셀의 기본 단락 제거
    p = container_box.add_paragraph()
    p.add_run('등급 배분').bold = True
    container_box.add_paragraph(f"• 배분 방식: {dist_method}")
    container_box.add_paragraph(f"• Process: {process_flow}")
    doc.add_paragraph()

    # --- 2단 레이아웃을 위한 메인 테이블 생성 ---
    # 3행 2열의 테이블을 만들고, 테두리는 보이지 않게 처리하여 레이아웃용으로만 사용
    layout_table = doc.add_table(rows=3, cols=2)
    layout_table.autofit = False
    layout_table.allow_autofit = False
    
    # 열 너비 설정 (A4용지 기준, 왼쪽:제목, 오른쪽:내용)
    layout_table.columns[0].width = Cm(4)
    layout_table.columns[1].width = Cm(13)

    # 섹션 데이터
    sections = {
        0: {"title": "등급별 분포 현황", "type": "table", "data": df},
        1: {"title": "구성원 VOE", "type": "list", "data": voe_list},
        2: {"title": "평가 운영상의 Issue", "type": "list", "data": issue_list},
    }

    for i, section in sections.items():
        # 왼쪽 셀 (제목)
        left_cell = layout_table.cell(i, 0)
        left_cell.text = ''  # 기본 단락 제거
        # 제목을 담을 테이블을 만들어 테두리 효과를 줌
        title_box = left_cell.add_table(rows=1, cols=1).cell(0,0)
        title_box.text = section["title"]
        title_box.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        left_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # 오른쪽 셀 (내용)
        right_cell = layout_table.cell(i, 1)
        right_cell.text = '' # 기본 단락 제거

        if section["type"] == "table":
            # 내용을 담을 테이블 추가 (셀 안에 테이블 추가)
            data_table = right_cell.add_table(rows=section["data"].shape[0] + 1, cols=section["data"].shape[1] + 1)
            data_table.style = 'Table Grid'
            # 헤더
            hdr_cells = data_table.rows[0].cells
            hdr_cells[0].text = '직급'
            for j, col_name in enumerate(section["data"].columns):
                hdr_cells[j+1].text = col_name
            # 내용
            for k, (index, row) in enumerate(section["data"].iterrows()):
                row_cells = data_table.rows[k+1].cells
                row_cells[0].text = index
                for l, value in enumerate(row):
                    row_cells[l+1].text = str(value)

        elif section["type"] == "list":
            for item in section["data"]:
                right_cell.add_paragraph(item)

    # 메모리에 문서를 저장하여 BytesIO 객체로 반환
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# --- Streamlit UI 구성 ---
st.title(title)
st.markdown("---")

# 1. 등급 배분
with st.container(border=True):
    # Streamlit에서는 제목을 내부에 넣는 것이 더 깔끔해 보입니다.
    st.subheader("등급 배분")
    st.markdown(f"**배분 방식:** {dist_method}")
    st.markdown(f"**Process:** {process_flow}")

st.write("") # 간격

# 각 섹션을 2단 레이아웃으로 표시
def display_section(title, content, content_type):
    cols = st.columns([0.25, 0.75], gap="medium") # 비율 조정
    with cols[0]:
        st.container(border=True).subheader(title)
    with cols[1]:
        if content_type == 'table':
            st.table(content)
        elif content_type == 'list':
            for item in content:
                st.markdown(item)
    st.write("") # 섹션 간 간격

# 2. 등급별 분포 현황
display_section("등급별 분포 현황", df, 'table')

# 3. 구성원 VOE
display_section("구성원 VOE", voe_list, 'list')

# 4. 평가 운영상의 Issue
display_section("평가 운영상의 Issue", issue_list, 'list')

st.markdown("---")

# --- 다운로드 버튼 ---
st.write("### 보고서 다운로드")
st.info("Word 파일은 이미지와 유사한 2단 레이아웃으로 생성됩니다.")
word_file = create_word_document()
st.download_button(
    label="📥 Word 파일로 다운로드 (레이아웃 적용)",
    data=word_file,
    file_name="성과관리_운영현황_보고서_레이아웃.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
